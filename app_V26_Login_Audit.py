
import os, re, sqlite3, json, base64, html, zipfile, tempfile, hashlib
from pathlib import Path
from datetime import date, datetime
from io import BytesIO

import streamlit as st
import pandas as pd
from PIL import Image
from dateutil.relativedelta import relativedelta
import qrcode

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import landscape, A4
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except Exception:
    arabic_reshaper = None
    get_display = None

APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
EXPORT_DIR = APP_DIR / "exports"
BG_DIR = APP_DIR / "assets" / "backgrounds"
UPLOAD_DIR = APP_DIR / "assets" / "uploads"
TEMPLATE_DIR = APP_DIR / "templates"
FONT_DIR = APP_DIR / "assets" / "fonts"
LOGO_DIR = APP_DIR / "assets" / "logos"
STAMP_DIR = APP_DIR / "assets" / "stamps"
SIGN_DIR = APP_DIR / "assets" / "signatures"
DB_FILE = DATA_DIR / "certificates.db"
EXCEL_FILE = DATA_DIR / "issued_certificates.xlsx"
USERS_FILE = DATA_DIR / "users.json"
ACTIVITY_EXCEL_FILE = DATA_DIR / "activity_log.xlsx"

for p in [DATA_DIR, EXPORT_DIR, BG_DIR, UPLOAD_DIR, TEMPLATE_DIR, FONT_DIR, LOGO_DIR, STAMP_DIR, SIGN_DIR]:
    p.mkdir(parents=True, exist_ok=True)

PAGE_W, PAGE_H = landscape(A4)  # 841.89 x 595.28 تقريباً
GREEN = "#4f8264"
BLACK = "#111111"

SYSTEM_FONT_OPTIONS = {
    "Tahoma": [r"C:\Windows\Fonts\tahoma.ttf"],
    "Arial": [r"C:\Windows\Fonts\arial.ttf", "/System/Library/Fonts/Supplemental/Arial.ttf", "/Library/Fonts/Arial.ttf"],
    "Calibri": [r"C:\Windows\Fonts\calibri.ttf"],
    "Times New Roman": [r"C:\Windows\Fonts\times.ttf", "/System/Library/Fonts/Supplemental/Times New Roman.ttf"],
    "DejaVu Sans": ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"],
    "FreeSans": ["/usr/share/fonts/truetype/freefont/FreeSans.ttf"],
}

FONT_ALIASES = {
    "Amiri Regular": ["Amiri-Regular.ttf", "Amiri Regular.ttf"],
    "Amiri Bold": ["Amiri-Bold.ttf", "Amiri Bold.ttf", "Amiri-Bold (1).ttf"],
    "Amiri Italic": ["Amiri-Italic.ttf"],
    "Amiri BoldItalic": ["Amiri-BoldItalic.ttf"],
    "Tajawal Regular": ["Tajawal-Regular.ttf"],
    "Tajawal Medium": ["Tajawal-Medium.ttf"],
    "Tajawal Bold": ["Tajawal-Bold.ttf", "Tajawal-Bold (1).ttf"],
    "Tajawal ExtraBold": ["Tajawal-ExtraBold.ttf"],
    "Tajawal Black": ["Tajawal-Black.ttf"],
    "Tajawal Light": ["Tajawal-Light.ttf"],
    "Tajawal ExtraLight": ["Tajawal-ExtraLight.ttf"],
    "Noto Naskh Arabic Regular": ["NotoNaskhArabic-Regular.ttf", "NotoNaskhArabic-VariableFont_wght.ttf"],
    "Noto Naskh Arabic Bold": ["NotoNaskhArabic-Bold.ttf", "NotoNaskhArabic-VariableFont_wght.ttf"],
    "Cairo Regular": ["Cairo-Regular.ttf", "Cairo-VariableFont_slnt,wght.ttf"],
    "Cairo Bold": ["Cairo-Bold.ttf", "Cairo-VariableFont_slnt,wght.ttf"],
    "Jomhuria Regular": ["Jomhuria-Regular.ttf"],
    "Cairoline": ["Cairoline.ttf"],
}

FONT_OPTIONS = SYSTEM_FONT_OPTIONS.copy()

st.set_page_config(page_title="منصة شهادات الخبرة V26 Login Audit", page_icon="🏆", layout="wide")


# ---------------------- أدوات عامة ----------------------
def scan_project_fonts():
    """يقرأ أي ملفات خطوط TTF/OTF داخل assets/fonts تلقائياً مع أسماء ودّية متوافقة مع القوالب."""
    fonts = {}
    if FONT_DIR.exists():
        for label, filenames in FONT_ALIASES.items():
            paths = []
            for fn in filenames:
                p = FONT_DIR / fn
                if p.exists():
                    paths.append(str(p))
            if paths:
                fonts[label] = paths

        for p in sorted(list(FONT_DIR.glob("*.ttf")) + list(FONT_DIR.glob("*.otf"))):
            label = p.stem.replace("-", " ").replace("_", " ")
            label = re.sub(r"\s+", " ", label).strip()
            if label not in fonts:
                fonts[label] = [str(p)]
    return fonts

def get_font_options():
    opts = {}
    opts.update(scan_project_fonts())
    opts.update(SYSTEM_FONT_OPTIONS)
    return opts

def available_font_labels():
    opts = get_font_options()
    preferred = [
        "Amiri Regular", "Amiri Bold",
        "Tajawal Regular", "Tajawal Medium", "Tajawal Bold", "Tajawal ExtraBold",
        "Noto Naskh Arabic Regular", "Noto Naskh Arabic Bold",
        "Cairo Regular", "Cairo Bold",
        "Jomhuria Regular",
        "DejaVu Sans",
    ]
    labels = []
    for label in preferred:
        if label in opts and any(os.path.exists(p) for p in opts[label]):
            labels.append(label)
    for label, paths in opts.items():
        if label not in labels and any(os.path.exists(p) for p in paths):
            labels.append(label)
    return labels or ["DejaVu Sans"]

def _font_path_from_candidates(candidates):
    for p in candidates:
        if p and os.path.exists(str(p)):
            return str(p)
    return None

def find_font(label="Arial"):
    """بحث ذكي عن الخط يمنع الرجوع إلى Helvetica قدر الإمكان."""
    opts = get_font_options()

    direct = _font_path_from_candidates(opts.get(label, []))
    if direct:
        return direct

    label_low = str(label or "").lower().replace("-", " ").replace("_", " ")

    for name, paths in opts.items():
        name_low = name.lower().replace("-", " ").replace("_", " ")
        if label_low and (label_low in name_low or name_low in label_low):
            found = _font_path_from_candidates(paths)
            if found:
                return found

    if "tajawal" in label_low:
        families = ["tajawal", "amiri", "noto", "cairo", "dejavu"]
    elif "amiri" in label_low:
        families = ["amiri", "tajawal", "noto", "cairo", "dejavu"]
    elif "noto" in label_low or "naskh" in label_low:
        families = ["noto", "amiri", "tajawal", "cairo", "dejavu"]
    elif "cairo" in label_low:
        families = ["cairo", "tajawal", "amiri", "noto", "dejavu"]
    elif "jomhuria" in label_low:
        families = ["jomhuria", "amiri", "tajawal", "noto", "dejavu"]
    else:
        families = ["tajawal", "amiri", "noto", "cairo", "jomhuria", "dejavu"]

    for family in families:
        for name, paths in opts.items():
            if family in name.lower():
                found = _font_path_from_candidates(paths)
                if found:
                    return found

    for fallback in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
    ]:
        if os.path.exists(fallback):
            return fallback

    return None

def _safe_reportlab_font_name(label, suffix=""):
    base = re.sub(r"[^A-Za-z0-9_]", "_", str(label or "ArabicFont"))
    return (base[:42] + suffix) or ("ArabicFont" + suffix)

def register_font(label="Arial"):
    """
    تسجيل خط ReportLab مع منع الرجوع إلى Helvetica؛ لأن Helvetica لا يدعم العربية.
    إذا فشل الخط المطلوب نستخدم أفضل خط عربي متاح.
    """
    regular_name = _safe_reportlab_font_name(label, "_R")
    bold_name = _safe_reportlab_font_name(label, "_B")

    path = find_font(label)
    if path:
        try:
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, path))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, path))
            return regular_name, bold_name
        except Exception:
            pass

    fallback_path = find_font("DejaVu Sans")
    if fallback_path:
        try:
            if "DejaVuSansArabicSafe" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont("DejaVuSansArabicSafe", fallback_path))
            return "DejaVuSansArabicSafe", "DejaVuSansArabicSafe"
        except Exception:
            pass

    return "Helvetica", "Helvetica-Bold"

def apply_cloud_safe_default_fonts():
    """
    يفرض خطوطاً عربية مستقرة على PDF.
    سبب ذلك: بعض الخطوط مثل Tajawal/Variable Fonts قد تعرض مربعات عند استخدام arabic_reshaper مع ReportLab.
    Amiri أكثر ثباتاً مع الحروف العربية المشكلة داخل PDF.
    """
    if st.session_state.get("_cloud_fonts_applied"):
        return

    available = available_font_labels()
    body_font = choose_font("Amiri Regular", "DejaVu Sans", available)
    bold_font = choose_font("Amiri Bold", body_font, available)
    name_font = choose_font("Amiri Bold", bold_font, available)

    for key in ["intro", "contribution", "issued", "legal", "date_label", "qr_label"]:
        if key in st.session_state.layout:
            st.session_state.layout[key]["font"] = body_font

    for key in ["heading", "period", "org", "issue_date", "certificate_no"]:
        if key in st.session_state.layout:
            st.session_state.layout[key]["font"] = bold_font

    if "name" in st.session_state.layout:
        st.session_state.layout["name"]["font"] = name_font

    st.session_state["_cloud_fonts_applied"] = True

def clean_text_for_pdf(text):
    """
    تنظيف محارف مخفية أو غير مدعومة تسبب ظهور مربعات داخل PDF.
    """
    text = "" if text is None else str(text)
    replacements = {
        "\u200b": "",
        "\u200c": "",
        "\u200d": "",
        "\u200e": "",
        "\u200f": "",
        "\ufeff": "",
        "\u061c": "",
        "\u00a0": " ",
        "□": "",
        "�": "",
    }
    for a, b in replacements.items():
        text = text.replace(a, b)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()

def rtl(text):
    text = clean_text_for_pdf(text)
    if arabic_reshaper and get_display:
        try:
            return get_display(arabic_reshaper.reshape(text))
        except Exception:
            return text
    return text

def fmt_date(x):
    if isinstance(x, (datetime, date)):
        return x.strftime("%Y/%m/%d")
    try:
        return pd.to_datetime(x).strftime("%Y/%m/%d")
    except Exception:
        return str(x)

def duration_ar(start, end):
    try:
        s = pd.to_datetime(start).date()
        e = pd.to_datetime(end).date()
        r = relativedelta(e, s)
        parts = []
        if r.years:
            parts.append(f"{r.years} سنة" if r.years == 1 else f"{r.years} سنوات")
        if r.months:
            parts.append(f"{r.months} شهر" if r.months == 1 else f"{r.months} أشهر")
        if r.days:
            parts.append(f"{r.days} يوم")
        return " و".join(parts) if parts else "يوم واحد"
    except Exception:
        return ""

def gender_terms(title):
    female = "السيدة" in title or "الآنسة" in title
    return {
        "participated": "شاركت" if female else "شارك",
        "contributed": "ساهمت" if female else "ساهم",
        "request_pronoun": "طلبها" if female else "طلبه",
    }

def backgrounds():
    return sorted([p for p in BG_DIR.iterdir() if p.suffix.lower() in [".jpg", ".jpeg", ".png"]])

def file_to_data_url(path):
    if not path or not os.path.exists(str(path)):
        return ""
    ext = Path(path).suffix.lower().replace(".", "")
    mime = "jpeg" if ext in ["jpg", "jpeg"] else "png"
    b64 = base64.b64encode(Path(path).read_bytes()).decode("utf-8")
    return f"data:image/{mime};base64,{b64}"

def save_upload(upload, prefix):
    if not upload:
        return ""
    ext = Path(upload.name).suffix.lower()
    p = UPLOAD_DIR / f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    p.write_bytes(upload.getbuffer())
    return str(p)

def make_qr_bytes(data):
    img = qrcode.make(data)
    bio = BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio

def qr_data_url(data):
    bio = make_qr_bytes(data)
    b64 = base64.b64encode(bio.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{b64}"

def hex_to_color(h):
    try:
        return colors.HexColor(h)
    except Exception:
        return colors.black

def wrap_by_chars(text, max_chars=95):
    words = str(text).split()
    lines, line = [], ""
    for w in words:
        trial = (line + " " + w).strip()
        if len(trial) <= max_chars:
            line = trial
        else:
            if line:
                lines.append(line)
            line = w
    if line:
        lines.append(line)
    return lines


# ---------------------- قاعدة البيانات والسجل ----------------------
def init_db():
    con = sqlite3.connect(DB_FILE)
    con.execute("""
        CREATE TABLE IF NOT EXISTS certificates (
            certificate_no TEXT PRIMARY KEY,
            name TEXT,
            gender_title TEXT,
            start_date TEXT,
            end_date TEXT,
            issue_date TEXT,
            organization TEXT,
            mode TEXT,
            pdf_path TEXT,
            created_at TEXT,
            issued_by TEXT
        )
    """)
    con.execute("""
        CREATE TABLE IF NOT EXISTS activity_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT,
            username TEXT,
            action TEXT,
            details TEXT
        )
    """)
    con.commit()
    cols = [r[1] for r in con.execute("PRAGMA table_info(certificates)").fetchall()]
    if "issued_by" not in cols:
        con.execute("ALTER TABLE certificates ADD COLUMN issued_by TEXT")
        con.commit()
    con.close()

def next_no(prefix="MNH"):
    init_db()
    year = datetime.now().year
    con = sqlite3.connect(DB_FILE)
    n = con.execute("SELECT COUNT(*) FROM certificates WHERE certificate_no LIKE ?", (f"{prefix}-{year}-%",)).fetchone()[0] + 1
    con.close()
    return f"{prefix}-{year}-{n:06d}"

def current_username():
    return st.session_state.get("username", "unknown")

def log_activity(action, details=""):
    init_db()
    username = current_username()
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    con = sqlite3.connect(DB_FILE)
    con.execute("INSERT INTO activity_log (created_at, username, action, details) VALUES (?,?,?,?)",
                (created_at, username, action, str(details)))
    con.commit()
    con.close()

    row = {"created_at": created_at, "username": username, "action": action, "details": str(details)}
    df = pd.read_excel(ACTIVITY_EXCEL_FILE) if ACTIVITY_EXCEL_FILE.exists() else pd.DataFrame()
    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
    df.to_excel(ACTIVITY_EXCEL_FILE, index=False)

def log(values, pdf_path):
    init_db()
    row = {
        "certificate_no": values["certificate_no"],
        "name": values["name"],
        "gender_title": values["gender_title"],
        "start_date": values["start_date"],
        "end_date": values["end_date"],
        "issue_date": values["issue_date"],
        "organization": values["organization_name"],
        "mode": values["mode"],
        "pdf_path": str(pdf_path),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "issued_by": current_username(),
    }
    con = sqlite3.connect(DB_FILE)
    con.execute("""
        INSERT OR REPLACE INTO certificates
        (certificate_no, name, gender_title, start_date, end_date, issue_date, organization, mode, pdf_path, created_at, issued_by)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)
    """, tuple(row.values()))
    con.commit()
    con.close()

    df = pd.read_excel(EXCEL_FILE) if EXCEL_FILE.exists() else pd.DataFrame()
    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
    df.to_excel(EXCEL_FILE, index=False)
    log_activity("ISSUE_CERTIFICATE", f'{row["certificate_no"]} - {row["name"]}')



# ---------------------- النصوص والقالب ----------------------
DEFAULT_TEXTS = {
    "heading": "إلى من يهمه الأمر",
    "intro": 'تشهد إدارة المركز التعليمي المؤقت، المنفذ ضمن مشروع "المناهل" الإنساني والتعليمي، وبالتعاون مع مديرية التربية، بأن {title}:',
    "name": "{name}",
    "period": "قد {participated} في تنفيذ وتقديم أنشطة وخدمات تعليمية ضمن المركز خلال الفترة الممتدة من: {start} لغاية {end}{duration}.",
    "contribution": "وقد {contributed} خلال هذه الفترة في دعم العملية التعليمية والأنشطة المرتبطة بها ضمن إطار المشروع الإنساني والتعليمي القائم آنذاك.",
    "issued": "وقد أصدرت هذه الشهادة بناءً على {request_pronoun} للاستفادة منها عند الحاجة.",
    "legal": "ونود التوضيح بأن هذه الشهادة تثبت المشاركة في الأنشطة والخدمات التعليمية ضمن مركز تعليمي مؤقت مرتبط بمشروع إنساني، ولا تعتبر إثباتاً لعلاقة عمل رسمية أو عقد توظيف، ولا يترتب عليها أي حقوق قانونية أو مالية أو تأمينية أو مطالبات وظيفية حالية أو مستقبلية تجاه الجهة المنظمة أو الجهات الداعمة أو المتعاونة.",
    "org": "{organization}",
    "date_label": "حررت بتاريخ:",
    "issue_date": "{issue_date}",
    "certificate_no": "No: {certificate_no}",
    "qr_label": "التحقق الإلكتروني",
}

ELEMENT_NAMES = {
    "heading": "العنوان",
    "intro": "المقدمة",
    "name": "الاسم",
    "period": "فقرة الفترة",
    "contribution": "فقرة المساهمة",
    "issued": "سبب الإصدار",
    "legal": "الفقرة القانونية",
    "org": "اسم الجهة",
    "date_label": "عبارة التاريخ",
    "issue_date": "التاريخ",
    "certificate_no": "رقم الشهادة",
    "qr": "QR",
    "qr_label": "عبارة QR",
    "logo": "الشعار",
    "stamp": "الختم",
    "signature": "التوقيع",
}


TEXT_KEYS = ["heading", "intro", "name", "period", "contribution", "issued", "legal"]
FOOTER_KEYS = ["org", "date_label", "issue_date", "certificate_no", "qr", "qr_label", "logo", "stamp", "signature"]
ALL_MOVE_KEYS = TEXT_KEYS + FOOTER_KEYS

def move_group(keys, dx=0, dy=0):
    for k in keys:
        if k in st.session_state.layout:
            st.session_state.layout[k]["x"] = float(st.session_state.layout[k].get("x", 0)) + dx
            st.session_state.layout[k]["y"] = float(st.session_state.layout[k].get("y", 0)) + dy

def scale_group_fonts(keys, delta=0):
    for k in keys:
        item = st.session_state.layout.get(k, {})
        if "size" in item:
            item["size"] = max(5, float(item.get("size", 10)) + delta)


DEFAULT_LAYOUT = {
    "heading": {"x":421, "y":350, "w":500, "size":18, "font":"Tahoma", "color":"#111111", "align":"center", "bold":True, "chars":80, "line_h":22, "visible":True},
    "intro": {"x":421, "y":316, "w":640, "size":12.5, "font":"Tahoma", "color":"#111111", "align":"center", "bold":False, "chars":105, "line_h":17, "visible":True},
    "name": {"x":421, "y":286, "w":420, "size":22, "font":"Tahoma", "color":GREEN, "align":"center", "bold":True, "chars":50, "line_h":25, "visible":True},
    "period": {"x":421, "y":254, "w":650, "size":12.5, "font":"Tahoma", "color":"#111111", "align":"center", "bold":True, "chars":95, "line_h":18, "visible":True},
    "contribution": {"x":421, "y":218, "w":660, "size":12.2, "font":"Tahoma", "color":"#111111", "align":"center", "bold":False, "chars":100, "line_h":17, "visible":True},
    "issued": {"x":421, "y":176, "w":620, "size":12.2, "font":"Tahoma", "color":"#111111", "align":"center", "bold":False, "chars":90, "line_h":17, "visible":True},
    "legal": {"x":421, "y":142, "w":660, "size":12, "font":"Tahoma", "color":"#111111", "align":"center", "bold":False, "chars":92, "line_h":17, "visible":True},
    "org": {"x":275, "y":78, "w":180, "size":12, "font":"Tahoma", "color":"#444444", "align":"center", "bold":True, "chars":30, "line_h":14, "visible":True},
    "date_label": {"x":545, "y":88, "w":120, "size":8, "font":"Tahoma", "color":"#666666", "align":"center", "bold":False, "chars":20, "line_h":10, "visible":True},
    "issue_date": {"x":545, "y":75, "w":120, "size":10, "font":"Tahoma", "color":"#111111", "align":"center", "bold":True, "chars":20, "line_h":12, "visible":True},
    "certificate_no": {"x":95, "y":440, "w":160, "size":8, "font":"Tahoma", "color":"#111111", "align":"left", "bold":True, "chars":30, "line_h":10, "visible":False},
    "qr": {"x":115, "y":70, "w":56, "h":56, "visible":False},
    "qr_label": {"x":143, "y":58, "w":100, "size":7, "font":"Tahoma", "color":"#111111", "align":"center", "bold":False, "chars":20, "line_h":9, "visible":False},
    "logo": {"x":720, "y":465, "w":55, "h":55, "visible":False},
    "stamp": {"x":388, "y":58, "w":75, "h":75, "visible":False},
    "signature": {"x":245, "y":85, "w":95, "h":40, "visible":False},
}


STYLE_PRESETS = {
    "🏆 Manahil Official": {
        "font_regular": "Amiri Regular",
        "font_bold": "Amiri Bold",
        "fallback": "DejaVu Sans",
        "updates": {
            "heading": {"size": 30, "y": 340, "color": "#111111", "bold": True, "line_h": 34, "w": 520},
            "intro": {"size": 13.8, "y": 302, "line_h": 18, "w": 660},
            "name": {"size": 32, "y": 268, "color": "#4f8264", "bold": True, "w": 500},
            "period": {"size": 14.2, "y": 226, "line_h": 21, "bold": True, "w": 680},
            "contribution": {"size": 13.4, "y": 184, "line_h": 20, "w": 680},
            "issued": {"size": 13.2, "y": 150, "line_h": 19, "w": 650},
            "legal": {"size": 10.5, "y": 120, "line_h": 15, "color": "#333333", "w": 690},
            "org": {"size": 12, "x": 275, "y": 78, "color": "#444444"},
            "date_label": {"size": 8, "x": 545, "y": 88},
            "issue_date": {"size": 10, "x": 545, "y": 75},
        },
    },
    "🎓 Academic": {
        "font_regular": "Noto Naskh Arabic Regular",
        "font_bold": "Noto Naskh Arabic Bold",
        "fallback": "Times New Roman",
        "updates": {
            "heading": {"size": 28, "y": 345, "color": "#111111"},
            "intro": {"size": 14, "y": 308, "line_h": 20},
            "name": {"size": 31, "y": 270, "color": "#3f6f56"},
            "period": {"size": 14, "y": 228, "line_h": 21},
            "contribution": {"size": 13, "y": 188, "line_h": 20},
            "issued": {"size": 13, "y": 154, "line_h": 19},
            "legal": {"size": 10.5, "y": 124, "line_h": 15, "color": "#333333"},
        },
    },
    "🏢 Corporate": {
        "font_regular": "Cairo Regular",
        "font_bold": "Cairo Bold",
        "fallback": "Arial",
        "updates": {
            "heading": {"size": 28, "y": 342, "color": "#222222"},
            "intro": {"size": 13, "y": 306, "line_h": 18},
            "name": {"size": 30, "y": 270, "color": "#4f8264"},
            "period": {"size": 13.5, "y": 230, "line_h": 20},
            "contribution": {"size": 12.8, "y": 192, "line_h": 19},
            "issued": {"size": 12.8, "y": 160, "line_h": 18},
            "legal": {"size": 10.2, "y": 130, "line_h": 15, "color": "#444444"},
        },
    },
    "📜 Classic": {
        "font_regular": "Amiri Regular",
        "font_bold": "Amiri Bold",
        "fallback": "Times New Roman",
        "updates": {
            "heading": {"size": 31, "y": 345, "color": "#111111"},
            "intro": {"size": 15, "y": 307, "line_h": 21},
            "name": {"size": 34, "y": 267, "color": "#4f8264"},
            "period": {"size": 15, "y": 224, "line_h": 22},
            "contribution": {"size": 14, "y": 180, "line_h": 21},
            "issued": {"size": 14, "y": 145, "line_h": 20},
            "legal": {"size": 11, "y": 112, "line_h": 16, "color": "#333333"},
        },
    },
}

def choose_font(preferred, fallback, available):
    if preferred in available:
        return preferred
    # جرّب إزالة Regular/Bold عند الحاجة
    base = preferred.replace(" Regular", "").replace(" Bold", "")
    for f in available:
        if base.lower() in f.lower():
            return f
    return fallback if fallback in available else (available[0] if available else "Arial")

def apply_style_preset(name, available):
    preset = STYLE_PRESETS.get(name)
    if not preset:
        return
    reg = choose_font(preset["font_regular"], preset["fallback"], available)
    bold = choose_font(preset["font_bold"], reg, available)
    for key, updates in preset["updates"].items():
        if key not in st.session_state.layout:
            continue
        st.session_state.layout[key].update(updates)
        if key in ["heading", "name", "period", "org", "issue_date"]:
            st.session_state.layout[key]["font"] = bold
            st.session_state.layout[key]["bold"] = True
        else:
            st.session_state.layout[key]["font"] = reg


def ensure_state():
    if "layout" not in st.session_state:
        st.session_state.layout = json.loads(json.dumps(DEFAULT_LAYOUT))
    if "texts_template" not in st.session_state:
        st.session_state.texts_template = DEFAULT_TEXTS.copy()
    apply_cloud_safe_default_fonts()

def render_texts(values):
    title = values.get("gender_title", "السيد")
    terms = gender_terms(title)
    dur = duration_ar(values.get("start_date", ""), values.get("end_date", ""))
    dur_part = f"، وذلك لمدة {dur}" if values.get("show_duration") and dur else ""
    fmt = {
        "name": values.get("name", ""),
        "title": title,
        "start": values.get("start_date", ""),
        "end": values.get("end_date", ""),
        "duration": dur_part,
        "participated": terms["participated"],
        "contributed": terms["contributed"],
        "request_pronoun": terms["request_pronoun"],
        "organization": values.get("organization_name", ""),
        "issue_date": values.get("issue_date", ""),
        "certificate_no": values.get("certificate_no", ""),
    }
    out = {}
    for k, t in st.session_state.texts_template.items():
        try:
            out[k] = t.format(**fmt)
        except Exception:
            out[k] = t
    return out


def auto_layout_manahil(values=None, available=None):
    """
    ترتيب ذكي قريب من النموذج الذي اعتمده المستخدم:
    يحسب أماكن العناصر الأساسية بشكل ثابت ومتوازن، مع بقاء إمكانية التعديل اليدوي بعد ذلك.
    """
    if available is None:
        available = available_font_labels()

    # اختيار أفضل خط متاح
    regular = choose_font("Amiri Regular", "DejaVu Sans", available)
    bold = choose_font("Amiri Bold", regular, available)

    L = st.session_state.layout

    # الخطوط العامة
    for k in ["intro", "contribution", "issued", "legal", "date_label", "qr_label"]:
        if k in L:
            L[k]["font"] = regular
            L[k]["bold"] = False
    for k in ["heading", "name", "period", "org", "issue_date", "certificate_no"]:
        if k in L:
            L[k]["font"] = bold
            L[k]["bold"] = True

    # مواقع وأحجام Auto Layout - مناهل رسمي
    updates = {
        # رأس الشهادة: الشعار موجود في الخلفية، لذلك نبدأ أسفله
        "certificate_no": {"x": 105, "y": 418, "size": 7.5, "w": 180, "align": "left", "color": "#111111"},

        "heading": {"x": 421, "y": 350, "w": 520, "size": 21, "line_h": 25, "chars": 60, "align": "center", "color": "#111111"},
        "intro": {"x": 421, "y": 308, "w": 665, "size": 11.2, "line_h": 15, "chars": 105, "align": "center", "color": "#111111"},

        "name": {"x": 421, "y": 262, "w": 520, "size": 25, "line_h": 30, "chars": 45, "align": "center", "color": "#4f8264"},

        "period": {"x": 421, "y": 218, "w": 690, "size": 11.8, "line_h": 16.5, "chars": 98, "align": "center", "color": "#111111"},
        "contribution": {"x": 421, "y": 182, "w": 690, "size": 10.8, "line_h": 15.5, "chars": 108, "align": "center", "color": "#111111"},
        "issued": {"x": 421, "y": 152, "w": 640, "size": 10.8, "line_h": 15.5, "chars": 92, "align": "center", "color": "#111111"},
        "legal": {"x": 421, "y": 122, "w": 700, "size": 9.6, "line_h": 13.2, "chars": 120, "align": "center", "color": "#222222"},

        # التذييل
        "qr": {"x": 118, "y": 72, "w": 45, "h": 45},
        "qr_label": {"x": 141, "y": 56, "w": 110, "size": 6.5, "line_h": 8, "align": "center", "chars": 20, "color": "#111111"},
        "org": {"x": 290, "y": 82, "w": 180, "size": 11.5, "line_h": 14, "align": "center", "chars": 25, "color": "#333333"},
        "date_label": {"x": 555, "y": 86, "w": 120, "size": 7, "line_h": 9, "align": "center", "chars": 20, "color": "#666666"},
        "issue_date": {"x": 555, "y": 72, "w": 120, "size": 9, "line_h": 11, "align": "center", "chars": 20, "color": "#111111"},

        # عناصر اختيارية
        "logo": {"x": 720, "y": 465, "w": 55, "h": 55},
        "stamp": {"x": 388, "y": 58, "w": 75, "h": 75},
        "signature": {"x": 245, "y": 85, "w": 95, "h": 40},
    }

    for key, vals in updates.items():
        if key in L:
            L[key].update(vals)

    # الإظهار حسب وضع الإخراج
    if values:
        L["certificate_no"]["visible"] = bool(values.get("show_no"))
        L["qr"]["visible"] = bool(values.get("show_qr"))
        L["qr_label"]["visible"] = bool(values.get("show_qr"))


def auto_layout_flow(values=None, available=None):
    """
    Auto Layout تدفقي: يحدد بداية النص ثم ينزل العناصر واحداً بعد الآخر حسب المسافات.
    مناسب إذا تغير طول الفقرات أو حجم الخط.
    """
    if available is None:
        available = available_font_labels()

    regular = choose_font("Amiri Regular", "DejaVu Sans", available)
    bold = choose_font("Amiri Bold", regular, available)
    L = st.session_state.layout

    # إعدادات من session_state أو افتراضية
    top_y = float(st.session_state.get("auto_top_y", 350))
    section_gap = float(st.session_state.get("auto_section_gap", 16))
    small_gap = float(st.session_state.get("auto_small_gap", 8))
    center_x = float(st.session_state.get("auto_center_x", 421))
    body_w = float(st.session_state.get("auto_body_w", 690))

    # أحجام
    sizes = {
        "heading": float(st.session_state.get("auto_heading_size", 21)),
        "intro": float(st.session_state.get("auto_intro_size", 11.2)),
        "name": float(st.session_state.get("auto_name_size", 25)),
        "period": float(st.session_state.get("auto_period_size", 11.8)),
        "body": float(st.session_state.get("auto_body_size", 10.8)),
        "legal": float(st.session_state.get("auto_legal_size", 9.6)),
    }

    text_templates = st.session_state.texts_template
    temp_values = values or {}
    texts = render_texts(temp_values) if temp_values else {k: text_templates.get(k, "") for k in text_templates}

    def estimate_height(key, chars, size, line_h):
        txt = texts.get(key, "")
        n_lines = max(1, len(wrap_by_chars(txt, int(chars))))
        return n_lines * line_h

    # تحديث خصائص أساسية
    base = {
        "heading": {"x": center_x, "w": 520, "size": sizes["heading"], "font": bold, "bold": True, "line_h": sizes["heading"]+4, "chars": 60},
        "intro": {"x": center_x, "w": body_w, "size": sizes["intro"], "font": regular, "bold": False, "line_h": sizes["intro"]+4, "chars": 108},
        "name": {"x": center_x, "w": 520, "size": sizes["name"], "font": bold, "bold": True, "line_h": sizes["name"]+5, "chars": 45, "color": "#4f8264"},
        "period": {"x": center_x, "w": body_w, "size": sizes["period"], "font": bold, "bold": True, "line_h": sizes["period"]+5, "chars": 100},
        "contribution": {"x": center_x, "w": body_w, "size": sizes["body"], "font": regular, "bold": False, "line_h": sizes["body"]+4.5, "chars": 112},
        "issued": {"x": center_x, "w": 650, "size": sizes["body"], "font": regular, "bold": False, "line_h": sizes["body"]+4.5, "chars": 95},
        "legal": {"x": center_x, "w": 700, "size": sizes["legal"], "font": regular, "bold": False, "line_h": sizes["legal"]+3.5, "chars": 122, "color": "#222222"},
    }

    y = top_y
    for key in ["heading", "intro", "name", "period", "contribution", "issued", "legal"]:
        L[key].update(base[key])
        L[key]["align"] = "center"
        L[key]["visible"] = True
        L[key]["y"] = y
        h = estimate_height(key, L[key].get("chars", 100), L[key].get("size", 11), L[key].get("line_h", 15))
        y -= h
        if key in ["heading", "intro", "name"]:
            y -= section_gap
        else:
            y -= small_gap

    # إذا نزل القانون كثيراً، صغّره قليلاً تلقائياً
    if L["legal"]["y"] < 95:
        L["legal"]["size"] = max(8.2, L["legal"]["size"] - 1.0)
        L["legal"]["line_h"] = max(11, L["legal"]["line_h"] - 1.0)
        L["legal"]["chars"] = 135

    # التذييل ثابت وهادئ
    footer = {
        "certificate_no": {"x": 105, "y": 418, "size": 7.5, "w": 180, "align": "left"},
        "qr": {"x": 118, "y": 72, "w": 45, "h": 45},
        "qr_label": {"x": 141, "y": 56, "w": 110, "size": 6.5, "line_h": 8, "align": "center"},
        "org": {"x": 290, "y": 82, "w": 180, "size": 11.5, "font": bold, "bold": True, "align": "center"},
        "date_label": {"x": 555, "y": 86, "w": 120, "size": 7, "font": regular, "bold": False, "align": "center"},
        "issue_date": {"x": 555, "y": 72, "w": 120, "size": 9, "font": bold, "bold": True, "align": "center"},
    }
    for key, vals in footer.items():
        if key in L:
            L[key].update(vals)

    if values:
        L["certificate_no"]["visible"] = bool(values.get("show_no"))
        L["qr"]["visible"] = bool(values.get("show_qr"))
        L["qr_label"]["visible"] = bool(values.get("show_qr"))



# ---------------------- معاينة المصمم ----------------------
def css_align_to_transform(align):
    if align == "center":
        return "translateX(-50%)"
    if align == "right":
        return "translateX(-100%)"
    return "none"

def preview_html(bg_path, texts, values, logo_path="", stamp_path="", sign_path="", scale=1.0):
    layout = st.session_state.layout
    bg_url = file_to_data_url(bg_path)
    qr_url = qr_data_url(f"Certificate No: {values.get('certificate_no','')}\nName: {values.get('name','')}\nIssue Date: {values.get('issue_date','')}")
    logo_url = file_to_data_url(logo_path)
    stamp_url = file_to_data_url(stamp_path)
    sign_url = file_to_data_url(sign_path)

    selected = st.session_state.get("selected_element", "heading")
    s = float(scale)
    width = PAGE_W * s
    height = PAGE_H * s

    bg_style = f"background-image:url('{bg_url}');" if bg_url else "background:#fff;"
    body = []
    for key in ["heading","intro","name","period","contribution","issued","legal","org","date_label","issue_date","certificate_no","qr_label"]:
        item = layout.get(key, {})
        if not item.get("visible", True):
            continue
        if key == "certificate_no" and not values.get("show_no"):
            continue
        if key == "qr_label" and not values.get("show_qr"):
            continue
        text = html.escape(clean_text_for_pdf(texts.get(key, "")))
        left = item.get("x", 0) * s
        top = (PAGE_H - item.get("y", 0)) * s
        w = item.get("w", 200) * s
        fs = item.get("size", 12) * s
        lh = item.get("line_h", 16) * s
        align = item.get("align", "center")
        weight = "700" if item.get("bold") else "400"
        border = "2px dashed #2f80ed" if key == selected else "1px dashed transparent"
        bgsel = "rgba(47,128,237,.05)" if key == selected else "transparent"
        transform = css_align_to_transform(align)
        body.append(f"""
        <div class="el" title="{ELEMENT_NAMES.get(key,key)}" style="
            left:{left}px; top:{top}px; width:{w}px; transform:{transform};
            font-family:'{item.get('font','Tahoma')}', Tahoma, Arial, sans-serif;
            font-size:{fs}px; line-height:{lh}px; color:{item.get('color','#111')};
            text-align:{align}; font-weight:{weight}; border:{border}; background:{bgsel};
            direction:rtl; white-space:normal;">
            {text}
        </div>""")

    for key, url in [("qr", qr_url if values.get("show_qr") else ""), ("logo", logo_url), ("stamp", stamp_url), ("signature", sign_url)]:
        item = layout.get(key, {})
        if not item.get("visible", False) or not url:
            continue
        left = item.get("x", 0) * s
        top = (PAGE_H - item.get("y", 0) - item.get("h", 50)) * s
        w = item.get("w", 50) * s
        h = item.get("h", 50) * s
        border = "2px dashed #2f80ed" if key == selected else "1px dashed transparent"
        body.append(f"""
        <img class="el img" title="{ELEMENT_NAMES.get(key,key)}" src="{url}"
            style="left:{left}px; top:{top}px; width:{w}px; height:{h}px; object-fit:contain; border:{border};" />
        """)

    html_doc = f"""
    <html><head><meta charset="utf-8">
    <style>
    body {{ margin:0; background:#f6f7f8; }}
    .wrap {{
        width:{width}px; height:{height}px; margin:0 auto; position:relative;
        {bg_style} background-size:100% 100%; background-repeat:no-repeat;
        border:5px solid #4f8264; border-radius:10px; box-shadow:0 8px 24px rgba(0,0,0,.12);
        overflow:hidden;
    }}
    .el {{ position:absolute; padding:2px 5px; box-sizing:border-box; }}
    .img {{ position:absolute; box-sizing:border-box; }}
    .hint {{
        width:{width}px; margin:10px auto 0; padding:10px 14px; border-radius:8px;
        background:#eef7f1; color:#2d5b42; font-family:Tahoma,Arial; direction:rtl; text-align:right; font-size:13px;
    }}
    </style></head>
    <body>
      <div class="wrap">{''.join(body)}</div>
      <div class="hint">المعاينة حية حسب الإعدادات. اختر العنصر من اللوحة اليسرى ثم عدّل الخط والحجم والموقع. إصدار PDF يستخدم نفس الإحداثيات.</div>
    </body></html>
    """
    return html_doc


# ---------------------- توليد PDF ----------------------
def draw_text_block(c, key, text):
    item = st.session_state.layout[key]
    if not item.get("visible", True):
        return
    font, bold = register_font(item.get("font", "Arial"))
    f = bold if item.get("bold") else font
    c.setFont(f, float(item.get("size", 12)))
    c.setFillColor(hex_to_color(item.get("color", "#111111")))
    x = float(item.get("x", PAGE_W/2))
    y = float(item.get("y", PAGE_H/2))
    align = item.get("align", "center")
    text = clean_text_for_pdf(text)
    lines = wrap_by_chars(text, int(item.get("chars", 90)))
    for line in lines:
        shaped = rtl(line)
        if align == "left":
            c.drawString(x, y, shaped)
        elif align == "right":
            c.drawRightString(x, y, shaped)
        else:
            c.drawCentredString(x, y, shaped)
        y -= float(item.get("line_h", 16))

def draw_image(c, path, key):
    item = st.session_state.layout[key]
    if not item.get("visible", False) or not path or not os.path.exists(str(path)):
        return
    c.drawImage(str(path), float(item["x"]), float(item["y"]), width=float(item["w"]), height=float(item["h"]), preserveAspectRatio=True, mask="auto")

def build_pdf(values, bg_path, logo_path="", stamp_path="", sign_path="", out_path="out.pdf"):
    c = canvas.Canvas(out_path, pagesize=landscape(A4))
    if bg_path and os.path.exists(str(bg_path)):
        c.drawImage(str(bg_path), 0, 0, width=PAGE_W, height=PAGE_H)
    else:
        c.setFillColor(colors.white)
        c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)

    texts = render_texts(values)

    keys = ["heading","intro","name","period","contribution","issued","legal","org","date_label","issue_date"]
    if values.get("show_no"):
        st.session_state.layout["certificate_no"]["visible"] = True
        keys.append("certificate_no")
    for k in keys:
        draw_text_block(c, k, texts.get(k, ""))

    if values.get("show_qr"):
        st.session_state.layout["qr"]["visible"] = True
        st.session_state.layout["qr_label"]["visible"] = True
        qr = make_qr_bytes(f"Certificate No: {values.get('certificate_no','')}\nName: {values.get('name','')}\nIssue Date: {values.get('issue_date','')}")
        q = st.session_state.layout["qr"]
        c.drawImage(ImageReader(qr), float(q["x"]), float(q["y"]), width=float(q["w"]), height=float(q["h"]))
        draw_text_block(c, "qr_label", texts.get("qr_label", ""))

    draw_image(c, logo_path, "logo")
    draw_image(c, stamp_path, "stamp")
    draw_image(c, sign_path, "signature")
    c.showPage()
    c.save()


# ---------------------- واجهة التحكم ----------------------
def nudge_selected(dx=0, dy=0):
    key = st.session_state.get("selected_element", "heading")
    item = st.session_state.layout[key]
    item["x"] = float(item.get("x", 0)) + dx
    item["y"] = float(item.get("y", 0)) + dy

def control_selected(font_labels):
    key = st.session_state.get("selected_element", "heading")
    item = st.session_state.layout[key]
    st.markdown(f"### ✏️ تعديل: {ELEMENT_NAMES.get(key,key)}")
    st.caption("تعديل سريع: استخدم الأسهم للتحريك مباشرة، أو غيّر X/Y والأحجام من الحقول أدناه.")

    # نص العنصر
    if key in DEFAULT_TEXTS and key not in ["qr"]:
        st.session_state.texts_template[key] = st.text_area(
            "النص",
            value=st.session_state.texts_template.get(key, ""),
            height=95 if key in ["legal","intro"] else 70,
            key=f"text_{key}"
        )
        st.caption("يمكن استخدام: {name}, {title}, {start}, {end}, {duration}, {participated}, {contributed}, {request_pronoun}, {organization}, {issue_date}, {certificate_no}")

    item["visible"] = st.checkbox("إظهار هذا العنصر", bool(item.get("visible", True)), key=f"vis_{key}")

    if key in ["qr","logo","stamp","signature"]:
        c1, c2 = st.columns(2)
        with c1:
            item["x"] = st.number_input("X", 0.0, float(PAGE_W), float(item.get("x", 0)), 1.0, key=f"x_{key}")
            item["w"] = st.number_input("العرض", 5.0, 400.0, float(item.get("w", 50)), 1.0, key=f"w_{key}")
        with c2:
            item["y"] = st.number_input("Y", 0.0, float(PAGE_H), float(item.get("y", 0)), 1.0, key=f"y_{key}")
            item["h"] = st.number_input("الارتفاع", 5.0, 400.0, float(item.get("h", 50)), 1.0, key=f"h_{key}")
    else:
        c1, c2 = st.columns(2)
        with c1:
            item["font"] = st.selectbox("نوع الخط", font_labels, index=font_labels.index(item.get("font")) if item.get("font") in font_labels else 0, key=f"font_{key}")
            item["size"] = st.number_input("حجم الخط", 6.0, 60.0, float(item.get("size", 12)), 0.5, key=f"size_{key}")
            item["color"] = st.color_picker("لون الخط", item.get("color", "#111111"), key=f"color_{key}")
        with c2:
            item["bold"] = st.checkbox("عريض Bold", bool(item.get("bold", False)), key=f"bold_{key}")
            item["align"] = st.radio("المحاذاة", ["right","center","left"], index=["right","center","left"].index(item.get("align","center")), horizontal=True, key=f"align_{key}")
            item["w"] = st.number_input("عرض الصندوق", 30.0, float(PAGE_W), float(item.get("w", 300)), 5.0, key=f"boxw_{key}")

        c3, c4 = st.columns(2)
        with c3:
            item["x"] = st.number_input("X أفقي", 0.0, float(PAGE_W), float(item.get("x", 0)), 1.0, key=f"x_{key}")
            item["chars"] = st.number_input("عرض السطر بالحروف", 10, 180, int(item.get("chars", 80)), 1, key=f"chars_{key}")
        with c4:
            item["y"] = st.number_input("Y عمودي", 0.0, float(PAGE_H), float(item.get("y", 0)), 1.0, key=f"y_{key}")
            item["line_h"] = st.number_input("المسافة بين السطور", 6.0, 50.0, float(item.get("line_h", 16)), 0.5, key=f"lh_{key}")

    st.markdown("#### تحريك سريع")
    a,b,c,d,e = st.columns(5)
    with b:
        if st.button("⬆️", key=f"up_{key}"):
            nudge_selected(0, 5)
            st.rerun()
    with d:
        if st.button("⬇️", key=f"down_{key}"):
            nudge_selected(0, -5)
            st.rerun()
    with a:
        if st.button("⬅️", key=f"left_{key}"):
            nudge_selected(-5, 0)
            st.rerun()
    with e:
        if st.button("➡️", key=f"right_{key}"):
            nudge_selected(5, 0)
            st.rerun()


def save_template_file(name):
    safe = re.sub(r"[^A-Za-z0-9_\-\u0600-\u06FF]", "_", name).strip("_") or "template"
    path = TEMPLATE_DIR / f"{safe}.json"
    data = {
        "layout": st.session_state.layout,
        "texts_template": st.session_state.texts_template,
        "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path

def load_template_file(path):
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    st.session_state.layout = data.get("layout", DEFAULT_LAYOUT)
    st.session_state.texts_template = data.get("texts_template", DEFAULT_TEXTS)



def render_pdf_first_page_png(pdf_path, zoom=2.0):
    """
    يحوّل أول صفحة من PDF إلى صورة PNG للمعاينة داخل Streamlit.
    هذا أفضل من iframe لأن بعض المتصفحات تمنع عرض PDF بصيغة data URL.
    يحتاج: pip install pymupdf
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(pdf_path))
        page = doc[0]
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        out_png = DATA_DIR / "_live_pdf_preview.png"
        pix.save(str(out_png))
        doc.close()
        return out_png, None
    except Exception as e:
        return None, str(e)

def pdf_iframe_html(pdf_path, height=720):
    """خيار احتياطي فقط؛ قد لا يعمل في بعض المتصفحات."""
    try:
        b64 = base64.b64encode(Path(pdf_path).read_bytes()).decode("utf-8")
        return f"""
        <iframe
            src="data:application/pdf;base64,{b64}#toolbar=0&navpanes=0&scrollbar=0"
            width="100%"
            height="{height}"
            style="border:1px solid #ddd; border-radius:10px; background:#f5f5f5;">
        </iframe>
        """
    except Exception as e:
        return f"<div style='color:red;direction:rtl'>تعذر عرض معاينة PDF: {html.escape(str(e))}</div>"



# ---------------------- تصدير Word وإصدار جماعي ----------------------
def safe_filename(name):
    name = str(name or "").strip()
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = re.sub(r"\s+", "_", name)
    return name or "certificate"

def normalize_gender(value):
    txt = str(value or "").strip().lower()
    if txt in ["أنثى", "انثى", "female", "f", "woman", "سيدة", "السيدة", "الآنسة", "انسة"]:
        return "السيدة"
    if txt in ["ذكر", "male", "m", "man", "سيد", "السيد"]:
        return "السيد"
    return "السيدة" if "ة" in str(value or "") else "السيد"

def set_paragraph_rtl(paragraph):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)
    except Exception:
        pass

def add_docx_para(doc, text, size=12, bold=False, color="000000", align="center", font_name="Arial", space_after=6):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(float(size))
    run.font.name = font_name
    try:
        clean = str(color).replace("#", "")
        run.font.color.rgb = RGBColor.from_string(clean if len(clean) == 6 else "000000")
    except Exception:
        pass
    return p

def build_docx(values, bg_path="", out_path="certificate.docx"):
    """
    يصدر نسخة Word قابلة للتعديل.
    ملاحظة: DOCX هنا قابل للتحرير ومقارب بصرياً، أما PDF فهو النسخة الرسمية المطابقة للخلفية والإحداثيات.
    """
    try:
        from docx import Document
        from docx.shared import Cm
        from docx.enum.section import WD_ORIENT
    except Exception as e:
        raise RuntimeError("تحتاج تثبيت python-docx: pip install python-docx") from e

    texts = render_texts(values)
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)

    # خلفية كصورة أعلى الصفحة تقريبياً وليست Watermark حقيقي
    if bg_path and os.path.exists(str(bg_path)):
        try:
            p = doc.add_paragraph()
            p.alignment = 1
            run = p.add_run()
            run.add_picture(str(bg_path), width=Cm(25.5))
            p.paragraph_format.space_after = Cm(0.1)
        except Exception:
            pass

    # النص القابل للتعديل
    add_docx_para(doc, texts.get("certificate_no", ""), size=8, bold=True, align="left", font_name="Arial", space_after=2) if values.get("show_no") else None
    add_docx_para(doc, texts.get("heading", ""), size=20, bold=True, font_name="Arial", space_after=8)
    add_docx_para(doc, texts.get("intro", ""), size=11.5, font_name="Arial", space_after=7)
    add_docx_para(doc, texts.get("name", ""), size=24, bold=True, color="4F8264", font_name="Arial", space_after=8)
    add_docx_para(doc, texts.get("period", ""), size=12.5, bold=True, font_name="Arial", space_after=6)
    add_docx_para(doc, texts.get("contribution", ""), size=11.5, font_name="Arial", space_after=5)
    add_docx_para(doc, texts.get("issued", ""), size=11.5, font_name="Arial", space_after=6)
    add_docx_para(doc, texts.get("legal", ""), size=10, font_name="Arial", space_after=10)

    # تذييل بسيط
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.cell(0,0).text = values.get("organization_name", "")
    table.cell(0,1).text = "حررت بتاريخ: " + values.get("issue_date", "")
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"

    doc.save(out_path)

def get_row_value(row, possible_names, default=""):
    lower_map = {str(k).strip().lower(): k for k in row.index}
    for name in possible_names:
        k = lower_map.get(str(name).strip().lower())
        if k is not None:
            val = row.get(k)
            if pd.notna(val):
                return val
    return default

def values_from_bulk_row(row, base_values, cert_no):
    name = get_row_value(row, ["name", "الاسم", "الاسم الكامل", "full_name", "Full Name"], base_values.get("name", ""))
    gender = get_row_value(row, ["gender", "الجنس", "الصفة", "title", "اللقب"], base_values.get("gender_title", "السيد"))
    start = get_row_value(row, ["start", "start_date", "تاريخ البداية", "البداية", "من"], base_values.get("start_date", ""))
    end = get_row_value(row, ["end", "end_date", "تاريخ النهاية", "النهاية", "إلى", "الى"], base_values.get("end_date", ""))
    issue = get_row_value(row, ["issue_date", "تاريخ الإصدار", "تاريخ الاصدار", "الإصدار"], base_values.get("issue_date", ""))
    org = get_row_value(row, ["organization", "org", "الجهة", "اسم الجهة"], base_values.get("organization_name", "المناهل"))

    vals = dict(base_values)
    vals.update({
        "certificate_no": cert_no,
        "name": str(name),
        "gender_title": normalize_gender(gender),
        "start_date": fmt_date(start),
        "end_date": fmt_date(end),
        "issue_date": fmt_date(issue),
        "organization_name": str(org),
    })
    return vals

def create_bulk_zip(df, base_values, bg_path, logo_path="", stamp_path="", sign_path="", export_pdf=True, export_docx=True):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bulk_dir = EXPORT_DIR / f"bulk_{ts}"
    bulk_dir.mkdir(parents=True, exist_ok=True)
    zip_path = EXPORT_DIR / f"bulk_certificates_{ts}.zip"

    files = []
    for i, row in df.iterrows():
        cert_no = next_no("MNH")
        vals = values_from_bulk_row(row, base_values, cert_no)
        filename_base = f"{safe_filename(vals.get('name'))}_{cert_no}"
        if export_pdf:
            pdf_path = bulk_dir / f"{filename_base}.pdf"
            build_pdf(vals, str(bg_path), logo_path, stamp_path, sign_path, str(pdf_path))
            log(vals, pdf_path)
            files.append(pdf_path)
        if export_docx:
            docx_path = bulk_dir / f"{filename_base}.docx"
            build_docx(vals, str(bg_path), str(docx_path))
            files.append(docx_path)
            if not export_pdf:
                log(vals, docx_path)

    merged_pdf = None
    pdf_files = [f for f in files if str(f).lower().endswith(".pdf")]
    if pdf_files:
        try:
            merged_pdf = bulk_dir / f"ALL_CERTIFICATES_{ts}.pdf"
            merge_pdf_files(pdf_files, merged_pdf)
            files.append(merged_pdf)
        except Exception:
            merged_pdf = None

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, f.name)
    return zip_path, files, merged_pdf

def create_sample_bulk_excel():
    sample = pd.DataFrame([
        {"الاسم": "غفران السيد", "الجنس": "أنثى", "تاريخ البداية": "2015/09/15", "تاريخ النهاية": "2024/06/15", "تاريخ الإصدار": fmt_date(date.today()), "الجهة": "المناهل"},
        {"الاسم": "محمد أحمد", "الجنس": "ذكر", "تاريخ البداية": "2018/01/01", "تاريخ النهاية": "2023/12/31", "تاريخ الإصدار": fmt_date(date.today()), "الجهة": "المناهل"},
    ])
    bio = BytesIO()
    sample.to_excel(bio, index=False)
    bio.seek(0)
    return bio.getvalue()



def print_pdf_html(pdf_path, height=520):
    """يعرض زر طباعة للـ PDF داخل iframe."""
    try:
        b64 = base64.b64encode(Path(pdf_path).read_bytes()).decode("utf-8")
        iframe_id = "pdfPrintFrame"
        return f"""
        <div style="direction:rtl;text-align:right;font-family:Tahoma,Arial;margin:8px 0;">
            <button onclick="
                var f=document.getElementById('{iframe_id}');
                if(f && f.contentWindow){{f.contentWindow.focus(); f.contentWindow.print();}}
            " style="
                background:#4f8264;color:white;border:0;border-radius:8px;
                padding:10px 18px;font-size:15px;cursor:pointer;">
                🖨️ طباعة PDF مباشرة
            </button>
            <div style="font-size:12px;color:#666;margin-top:6px;">
                إذا لم تعمل الطباعة مباشرة، افتح الملف من زر التحميل واطبعه من المتصفح.
            </div>
        </div>
        <iframe id="{iframe_id}"
            src="data:application/pdf;base64,{b64}"
            width="100%" height="{height}"
            style="border:1px solid #ddd;border-radius:10px;background:#f5f5f5;">
        </iframe>
        """
    except Exception as e:
        return f"<div style='color:red;direction:rtl'>تعذر تجهيز الطباعة: {html.escape(str(e))}</div>"


def print_png_html(png_path, height=560):
    """يعرض صورة الصفحة الأولى ويطبعها مباشرة؛ أكثر ثباتاً من طباعة PDF داخل iframe."""
    try:
        b64 = base64.b64encode(Path(png_path).read_bytes()).decode("utf-8")
        return f"""
        <html>
        <head>
        <meta charset="utf-8">
        <style>
            body {{ margin:0; padding:0; background:#f5f5f5; font-family:Tahoma,Arial; }}
            .toolbar {{ direction:rtl; text-align:right; padding:10px; }}
            button {{
                background:#4f8264;color:white;border:0;border-radius:8px;
                padding:10px 18px;font-size:15px;cursor:pointer;
            }}
            .note {{font-size:12px;color:#666;margin-top:6px;direction:rtl;}}
            .page {{ text-align:center; }}
            img {{ max-width:100%; height:auto; border:1px solid #ddd; background:white; }}
            @media print {{
                .toolbar {{ display:none; }}
                body {{ background:white; }}
                img {{ width:100%; border:0; }}
                @page {{ size: A4 landscape; margin: 0; }}
            }}
        </style>
        </head>
        <body>
            <div class="toolbar">
                <button onclick="window.print()">🖨️ طباعة الصورة المطابقة للـ PDF</button>
                <div class="note">هذه الطباعة تعتمد على صورة مولّدة من ملف PDF الحقيقي، لذلك لا تظهر فارغة.</div>
            </div>
            <div class="page">
                <img src="data:image/png;base64,{b64}" />
            </div>
        </body>
        </html>
        """
    except Exception as e:
        return f"<div style='color:red;direction:rtl'>تعذر تجهيز صورة الطباعة: {html.escape(str(e))}</div>"

def print_preview_component(pdf_path, height=680):
    """يحاول تحويل PDF إلى صورة للطباعة، وإن فشل يعود لطريقة PDF القديمة."""
    png_path, err = render_pdf_first_page_png(pdf_path, zoom=2.4)
    if png_path and png_path.exists():
        return print_png_html(png_path, height=height)
    return print_pdf_html(pdf_path, height=height)

def merge_pdf_files(pdf_files, out_path):
    """يدمج ملفات PDF في ملف واحد للطباعة الجماعية. يعتمد على PyMuPDF."""
    try:
        import fitz
    except Exception as e:
        raise RuntimeError("تحتاج تثبيت pymupdf: pip install pymupdf") from e

    merged = fitz.open()
    for pdf in pdf_files:
        if pdf and Path(pdf).exists():
            doc = fitz.open(str(pdf))
            merged.insert_pdf(doc)
            doc.close()
    if merged.page_count == 0:
        raise RuntimeError("لا توجد ملفات PDF صالحة للدمج.")
    merged.save(str(out_path))
    merged.close()
    return out_path


# ---------------------- تسجيل الدخول والمستخدمون ----------------------
def hash_password(password):
    return hashlib.sha256(str(password).encode("utf-8")).hexdigest()

def ensure_users_file():
    if USERS_FILE.exists():
        return
    default_users = {
        "admin": {
            "password_hash": hash_password("admin123"),
            "role": "admin",
            "full_name": "System Admin",
            "active": True,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    }
    USERS_FILE.write_text(json.dumps(default_users, ensure_ascii=False, indent=2), encoding="utf-8")

def load_users():
    ensure_users_file()
    try:
        return json.loads(USERS_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}

def save_users(users):
    USERS_FILE.write_text(json.dumps(users, ensure_ascii=False, indent=2), encoding="utf-8")

def authenticate(username, password):
    users = load_users()
    user = users.get(str(username).strip())
    if not user or not user.get("active", True):
        return False
    return user.get("password_hash") == hash_password(password)

def get_user_role(username=None):
    username = username or current_username()
    return load_users().get(username, {}).get("role", "user")

def is_admin():
    return get_user_role() == "admin"

def login_screen():
    st.markdown("""
    <div style="max-width:520px;margin:35px auto;padding:28px;border:1px solid #d9eadf;border-radius:18px;background:#f8fbf9;direction:rtl;text-align:center">
      <h1 style="color:#2f5f46">🏆 منصة شهادات الخبرة</h1>
      <p>يرجى تسجيل الدخول لمتابعة إصدار الشهادات</p>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        username = st.text_input("اسم المستخدم")
        password = st.text_input("كلمة المرور", type="password")
        submitted = st.form_submit_button("تسجيل الدخول", use_container_width=True)
    if submitted:
        if authenticate(username, password):
            st.session_state["authenticated"] = True
            st.session_state["username"] = username.strip()
            log_activity("LOGIN", "User logged in")
            st.rerun()
        else:
            st.error("اسم المستخدم أو كلمة المرور غير صحيحة")
    st.info("المستخدم الافتراضي لأول تشغيل: admin / admin123 — غيّره فوراً من تبويب المستخدمين.")

def require_login():
    ensure_users_file()
    if not st.session_state.get("authenticated"):
        login_screen()
        st.stop()

def logout_button():
    with st.sidebar:
        st.markdown("### 👤 المستخدم")
        st.write(f"**{current_username()}**")
        st.caption(f"الصلاحية: {get_user_role()}")
        if st.button("تسجيل الخروج", use_container_width=True):
            log_activity("LOGOUT", "User logged out")
            st.session_state.clear()
            st.rerun()

def users_management_tab():
    st.subheader("👤 إدارة المستخدمين")
    if not is_admin():
        st.warning("هذه الصفحة متاحة للمدير فقط.")
        return
    users = load_users()
    rows = [{"username": u, "full_name": i.get("full_name",""), "role": i.get("role","user"), "active": i.get("active", True), "created_at": i.get("created_at","")} for u,i in users.items()]
    st.dataframe(pd.DataFrame(rows), use_container_width=True)

    st.divider()
    st.markdown("### إضافة / تحديث مستخدم")
    with st.form("add_user_form"):
        new_username = st.text_input("اسم المستخدم")
        full_name = st.text_input("الاسم الكامل")
        role = st.selectbox("الصلاحية", ["user", "admin", "viewer"])
        new_password = st.text_input("كلمة المرور", type="password")
        active = st.checkbox("فعّال", value=True)
        save_btn = st.form_submit_button("حفظ", use_container_width=True)
    if save_btn:
        if not new_username.strip() or not new_password:
            st.error("يرجى إدخال اسم المستخدم وكلمة المرور.")
        else:
            users[new_username.strip()] = {
                "password_hash": hash_password(new_password),
                "role": role,
                "full_name": full_name,
                "active": active,
                "created_at": users.get(new_username.strip(), {}).get("created_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            }
            save_users(users)
            log_activity("UPSERT_USER", new_username.strip())
            st.success("تم حفظ المستخدم.")
            st.rerun()

    st.divider()
    st.markdown("### تغيير كلمة مرور المستخدم الحالي")
    with st.form("change_password_form"):
        old_pass = st.text_input("كلمة المرور الحالية", type="password")
        new_pass = st.text_input("كلمة المرور الجديدة", type="password")
        confirm_pass = st.text_input("تأكيد كلمة المرور", type="password")
        change_btn = st.form_submit_button("تغيير كلمة المرور", use_container_width=True)
    if change_btn:
        username = current_username()
        if not authenticate(username, old_pass):
            st.error("كلمة المرور الحالية غير صحيحة.")
        elif new_pass != confirm_pass:
            st.error("كلمة المرور غير متطابقة.")
        elif len(new_pass) < 6:
            st.error("كلمة المرور يجب أن تكون 6 أحرف على الأقل.")
        else:
            users[username]["password_hash"] = hash_password(new_pass)
            save_users(users)
            log_activity("CHANGE_PASSWORD", username)
            st.success("تم تغيير كلمة المرور.")

def activity_log_tab():
    st.subheader("📋 سجل النشاطات")
    if ACTIVITY_EXCEL_FILE.exists():
        df = pd.read_excel(ACTIVITY_EXCEL_FILE)
        st.dataframe(df.tail(300).sort_index(ascending=False), use_container_width=True)
        st.download_button("تحميل سجل النشاطات Excel", ACTIVITY_EXCEL_FILE.read_bytes(), file_name="activity_log.xlsx")
    else:
        st.info("لا يوجد سجل نشاطات بعد.")


def main():
    ensure_state()
    init_db()
    require_login()
    logout_button()

    st.markdown("""
    <style>
    .block-container {padding-top: 1.1rem; max-width: 98%;}
    [data-testid="stMetricValue"] {font-size: 20px;}
    /* جعل لوحة التعديل أسهل للاستخدام وتقليل المسافات */
    div[data-testid="stVerticalBlock"] { gap: 0.55rem; }
    .stNumberInput input { text-align: center; }
    .designer-tip {
        background:#eef7f1; border:1px solid #cfe7d8; color:#24533b;
        padding:10px 12px; border-radius:10px; direction:rtl; text-align:right;
        font-family:Tahoma,Arial; margin-bottom:8px;
    }
    </style>
    """, unsafe_allow_html=True)

    st.title("🏆 منصة شهادات الخبرة - V26 Login & Audit")
    st.caption("نسخة مستقرة للعربية: اعتماد Amiri كخط PDF أساسي + تنظيف المحارف المخفية + إصلاح ظهور المربعات.")

    tab_issue, tab_bg, tab_templates, tab_log, tab_users, tab_activity = st.tabs(["🧾 إصدار شهادة", "🖼️ الخلفيات", "💾 القوالب", "📚 السجل", "👤 المستخدمون", "📋 النشاطات"])

    with tab_issue:
        font_labels = available_font_labels()
        bgs = backgrounds()

        left, right = st.columns([0.30, 0.70], gap="large")

        with left:
            st.subheader("بيانات الشهادة")
            if bgs:
                bg_names = [p.name for p in bgs]
                idx = bg_names.index("manahil_blank.jpg") if "manahil_blank.jpg" in bg_names else 0
                bg_path = st.selectbox("خلفية الشهادة", bgs, format_func=lambda p: p.name, index=idx)
            else:
                bg_path = ""
                st.warning("لا توجد خلفيات. ارفع خلفية من تبويب الخلفيات.")

            name = st.text_input("الاسم الكامل", "غفران السيد")
            gender_title = st.radio("الصفة", ["السيد", "السيدة", "الآنسة"], horizontal=True, index=1)
            start_date = st.date_input("تاريخ البداية", date(2015,9,15))
            end_date = st.date_input("تاريخ النهاية", date(2024,6,15))
            issue_date = st.date_input("تاريخ الإصدار", date.today())
            org = st.text_input("اسم الجهة أسفل الشهادة", "المناهل")

            mode = st.radio("وضع الإخراج", ["مطابق للنموذج الأصلي", "نسخة إلكترونية مع QR ورقم شهادة"], index=0)
            show_qr = mode.startswith("نسخة إلكترونية")
            show_no = mode.startswith("نسخة إلكترونية")
            show_duration = st.checkbox("إظهار مدة الفترة", value=False)
            cert_no = st.text_input("رقم الشهادة", next_no("MNH"))

            st.session_state.layout["certificate_no"]["visible"] = show_no
            st.session_state.layout["qr"]["visible"] = show_qr
            st.session_state.layout["qr_label"]["visible"] = show_qr

            values = {
                "certificate_no": cert_no,
                "name": name,
                "gender_title": gender_title,
                "start_date": fmt_date(start_date),
                "end_date": fmt_date(end_date),
                "issue_date": fmt_date(issue_date),
                "organization_name": org,
                "mode": mode,
                "show_qr": show_qr,
                "show_no": show_no,
                "show_duration": show_duration,
            }

            with st.expander("📎 ملفات اختيارية", expanded=False):
                logo = st.file_uploader("شعار اختياري", type=["png","jpg","jpeg"])
                stamp = st.file_uploader("ختم اختياري", type=["png","jpg","jpeg"])
                sign = st.file_uploader("توقيع اختياري", type=["png","jpg","jpeg"])
            # حفظ مؤقت للمعاينة بعد الرفع
            if logo:
                logo_path = save_upload(logo, "logo_preview")
                st.session_state.last_logo = logo_path
                st.session_state.layout["logo"]["visible"] = True
            else:
                logo_path = st.session_state.get("last_logo", "")
            if stamp:
                stamp_path = save_upload(stamp, "stamp_preview")
                st.session_state.last_stamp = stamp_path
                st.session_state.layout["stamp"]["visible"] = True
            else:
                stamp_path = st.session_state.get("last_stamp", "")
            if sign:
                sign_path = save_upload(sign, "sign_preview")
                st.session_state.last_sign = sign_path
                st.session_state.layout["signature"]["visible"] = True
            else:
                sign_path = st.session_state.get("last_sign", "")

            st.divider()
            st.subheader("🎨 Designer")
            st.markdown('<div class="designer-tip">اختر العنصر من هنا وعدّل فوراً، والمعاينة تبقى ظاهرة على يمين الصفحة. لا تحتاج للنزول إلى آخر الصفحة ثم الرجوع للمعاينة.</div>', unsafe_allow_html=True)
            preset_choice = st.selectbox("تنسيق جاهز سريع", ["— اختر —"] + list(STYLE_PRESETS.keys()))
            if preset_choice != "— اختر —" and st.button("تطبيق التنسيق الجاهز", use_container_width=True):
                apply_style_preset(preset_choice, font_labels)
                st.success(f"تم تطبيق التنسيق: {preset_choice}")
                st.rerun()
            with st.expander("🎯 التحريك الجماعي", expanded=True):
                st.caption("استخدم هذا القسم لتحريك مجموعة كاملة بدلاً من تعديل كل عنصر وحده.")
                group_choice = st.radio(
                    "المجموعة",
                    ["النصوص الرئيسية فقط", "التذييل والـ QR فقط", "كل العناصر"],
                    horizontal=False,
                    key="group_choice"
                )
                if group_choice == "النصوص الرئيسية فقط":
                    group_keys = TEXT_KEYS
                elif group_choice == "التذييل والـ QR فقط":
                    group_keys = FOOTER_KEYS
                else:
                    group_keys = ALL_MOVE_KEYS

                step = st.number_input("مقدار التحريك", 1.0, 100.0, 10.0, 1.0, key="group_step")

                g1, g2, g3, g4 = st.columns(4)
                with g1:
                    if st.button("⬆️ رفع المجموعة", use_container_width=True):
                        move_group(group_keys, 0, step)
                        st.rerun()
                with g2:
                    if st.button("⬇️ إنزال المجموعة", use_container_width=True):
                        move_group(group_keys, 0, -step)
                        st.rerun()
                with g3:
                    if st.button("➡️ يمين", use_container_width=True):
                        move_group(group_keys, step, 0)
                        st.rerun()
                with g4:
                    if st.button("⬅️ يسار", use_container_width=True):
                        move_group(group_keys, -step, 0)
                        st.rerun()

                st.markdown("**تكبير/تصغير خطوط النصوص دفعة واحدة**")
                f1, f2 = st.columns(2)
                with f1:
                    if st.button("A+ تكبير النصوص", use_container_width=True):
                        scale_group_fonts(TEXT_KEYS, 1)
                        st.rerun()
                with f2:
                    if st.button("A- تصغير النصوص", use_container_width=True):
                        scale_group_fonts(TEXT_KEYS, -1)
                        st.rerun()

            with st.expander("⚙️ الضبط التلقائي الذكي", expanded=False):
                st.caption("يرتب النصوص تلقائياً مثل النموذج المعتمد، ثم تستطيع تعديل أي عنصر يدوياً.")
                a1, a2 = st.columns(2)
                with a1:
                    st.session_state.auto_top_y = st.number_input("بداية النص Y", 280.0, 420.0, float(st.session_state.get("auto_top_y", 350.0)), 1.0)
                    st.session_state.auto_heading_size = st.number_input("حجم العنوان الذكي", 12.0, 36.0, float(st.session_state.get("auto_heading_size", 21.0)), 0.5)
                    st.session_state.auto_name_size = st.number_input("حجم الاسم الذكي", 14.0, 42.0, float(st.session_state.get("auto_name_size", 25.0)), 0.5)
                    st.session_state.auto_section_gap = st.number_input("المسافة بين الأقسام", 4.0, 35.0, float(st.session_state.get("auto_section_gap", 16.0)), 0.5)
                with a2:
                    st.session_state.auto_center_x = st.number_input("محور النص X", 250.0, 600.0, float(st.session_state.get("auto_center_x", 421.0)), 1.0)
                    st.session_state.auto_intro_size = st.number_input("حجم المقدمة الذكي", 8.0, 24.0, float(st.session_state.get("auto_intro_size", 11.2)), 0.5)
                    st.session_state.auto_body_size = st.number_input("حجم المتن الذكي", 8.0, 24.0, float(st.session_state.get("auto_body_size", 10.8)), 0.5)
                    st.session_state.auto_legal_size = st.number_input("حجم القانون الذكي", 7.0, 18.0, float(st.session_state.get("auto_legal_size", 9.6)), 0.5)
                b1, b2 = st.columns(2)
                with b1:
                    if st.button("🔄 تطبيق ترتيب مناهل الرسمي", use_container_width=True):
                        auto_layout_manahil(values, font_labels)
                        st.success("تم تطبيق ترتيب مناهل الرسمي.")
                        st.rerun()
                with b2:
                    if st.button("🧠 تطبيق ترتيب تدفقي ذكي", use_container_width=True):
                        auto_layout_flow(values, font_labels)
                        st.success("تم تطبيق الترتيب الذكي.")
                        st.rerun()

            st.session_state.selected_element = st.selectbox(
                "اختر العنصر المراد تعديله",
                list(ELEMENT_NAMES.keys()),
                format_func=lambda k: ELEMENT_NAMES[k],
                index=list(ELEMENT_NAMES.keys()).index(st.session_state.get("selected_element","heading"))
            )
            # شريط تعديل سريع قريب من اختيار العنصر
            _key = st.session_state.get("selected_element", "heading")
            st.markdown("**تحريك سريع قريب من المعاينة**")
            m1, m2, m3, m4, m5 = st.columns(5)
            with m1:
                if st.button("⬅️ -10", key="top_left"):
                    nudge_selected(-10, 0); st.rerun()
            with m2:
                if st.button("⬆️ +10", key="top_up"):
                    nudge_selected(0, 10); st.rerun()
            with m3:
                if st.button("⬇️ -10", key="top_down"):
                    nudge_selected(0, -10); st.rerun()
            with m4:
                if st.button("➡️ +10", key="top_right"):
                    nudge_selected(10, 0); st.rerun()
            with m5:
                if st.button("🎯 وسط", key="top_center"):
                    st.session_state.layout[_key]["x"] = 421
                    st.rerun()

            control_selected(font_labels)

            st.divider()
            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button("↩️ إعادة ضبط القالب"):
                    st.session_state.layout = json.loads(json.dumps(DEFAULT_LAYOUT))
                    st.session_state.texts_template = DEFAULT_TEXTS.copy()
                    st.rerun()
            with c2:
                tname_quick = st.text_input("اسم القالب", "manahil_designer")
            with c3:
                if st.button("💾 حفظ كقالب"):
                    p = save_template_file(tname_quick)
                    st.success(f"تم حفظ القالب: {p.name}")

            st.divider()
            pdf_col1, pdf_col2 = st.columns(2)
            with pdf_col1:
                preview_scale = st.slider("حجم المعاينة", 0.65, 1.15, 0.90, 0.05)
                exact_pdf_preview = st.checkbox("معاينة PDF مطابقة للتحميل", value=True)
            with pdf_col2:
                issue_btn = st.button("🖨️ إصدار الشهادة PDF", type="primary", use_container_width=True)

        with right:
            texts = render_texts(values)
            st.subheader("معاينة حية للشهادة")
            st.caption("نصيحة: فعّل معاينة PDF المطابقة حتى ترى نفس نتيجة الملف المحمّل، وليس مجرد معاينة HTML تقريبية.")
            if exact_pdf_preview:
                # هذه المعاينة تستخدم نفس دالة إصدار PDF ثم تحول الصفحة إلى صورة، لذلك لا تظهر فارغة.
                temp_preview = DATA_DIR / "_live_pdf_preview.pdf"
                try:
                    build_pdf(values, str(bg_path), logo_path, stamp_path, sign_path, str(temp_preview))
                    png_path, err = render_pdf_first_page_png(temp_preview, zoom=2.0)
                    if png_path and png_path.exists():
                        st.image(str(png_path), use_container_width=True)
                        st.caption("هذه صورة مولّدة من PDF الحقيقي، لذلك هي أقرب نتيجة لما سيتم تحميله.")
                    else:
                        st.warning("لم أستطع تحويل PDF إلى صورة. ثبّت PyMuPDF بالأمر: pip install pymupdf")
                        st.components.v1.html(pdf_iframe_html(temp_preview, height=720), height=740, scrolling=False)
                        if err:
                            st.caption(f"تفاصيل الخطأ: {err}")
                except Exception as e:
                    st.error(f"تعذر إنشاء معاينة PDF: {e}")
                    st.components.v1.html(
                        preview_html(bg_path, texts, values, logo_path, stamp_path, sign_path, preview_scale),
                        height=int(PAGE_H*preview_scale)+95,
                        scrolling=True
                    )
            else:
                st.components.v1.html(
                    preview_html(bg_path, texts, values, logo_path, stamp_path, sign_path, preview_scale),
                    height=int(PAGE_H*preview_scale)+95,
                    scrolling=True
                )

            if issue_btn:
                safe = re.sub(r"[^A-Za-z0-9_-]", "_", cert_no)
                out = EXPORT_DIR / f"{safe}.pdf"
                build_pdf(values, str(bg_path), logo_path, stamp_path, sign_path, str(out))
                log(values, out)
                st.success("تم إصدار الشهادة PDF بنجاح")
                st.download_button("⬇️ تحميل PDF", out.read_bytes(), file_name=out.name, mime="application/pdf")
                with st.expander("🖨️ طباعة الشهادة", expanded=True):
                    st.components.v1.html(print_preview_component(out, height=560), height=680, scrolling=True)

            st.divider()
            st.subheader("📝 تصدير Word")
            st.caption("Word قابل للتعديل ومناسب للمراجعة، أما PDF فهو النسخة الرسمية المطابقة للتصميم.")
            if st.button("📝 إنشاء نسخة Word DOCX", use_container_width=True):
                try:
                    safe = re.sub(r"[^A-Za-z0-9_-]", "_", cert_no)
                    docx_out = EXPORT_DIR / f"{safe}.docx"
                    build_docx(values, str(bg_path), str(docx_out))
                    st.success("تم إنشاء ملف Word بنجاح")
                    st.download_button("⬇️ تحميل Word DOCX", docx_out.read_bytes(), file_name=docx_out.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(str(e))
                    st.info("ثبّت مكتبة Word بالأمر: pip install python-docx")

            st.divider()
            st.subheader("📦 إصدار جماعي من Excel")
            st.caption("الأعمدة المقبولة: الاسم، الجنس، تاريخ البداية، تاريخ النهاية، تاريخ الإصدار، الجهة. ويمكن أيضاً استخدام أسماء إنجليزية مثل name, gender, start_date, end_date.")
            st.download_button("⬇️ تحميل نموذج Excel", create_sample_bulk_excel(), file_name="bulk_certificates_template.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            bulk_file = st.file_uploader("رفع ملف Excel للإصدار الجماعي", type=["xlsx"], key="bulk_xlsx")
            bpdf, bdocx = st.columns(2)
            with bpdf:
                bulk_pdf = st.checkbox("إصدار PDF جماعي", value=True)
            with bdocx:
                bulk_docx = st.checkbox("إصدار Word جماعي", value=False)
            if bulk_file:
                try:
                    bulk_df = pd.read_excel(bulk_file)
                    st.dataframe(bulk_df.head(20), use_container_width=True)
                    if st.button("🚀 إصدار جماعي ZIP", type="primary", use_container_width=True):
                        zip_out, files, merged_pdf = create_bulk_zip(bulk_df, values, bg_path, logo_path, stamp_path, sign_path, bulk_pdf, bulk_docx)
                        st.success(f"تم إنشاء {len(files)} ملف داخل ZIP")
                        st.download_button("⬇️ تحميل ZIP", zip_out.read_bytes(), file_name=zip_out.name, mime="application/zip")
                        if merged_pdf and merged_pdf.exists():
                            st.download_button("📚 تحميل PDF موحد للطباعة الجماعية", merged_pdf.read_bytes(), file_name=merged_pdf.name, mime="application/pdf")
                            with st.expander("🖨️ طباعة جميع الشهادات", expanded=True):
                                st.components.v1.html(print_preview_component(merged_pdf, height=600), height=720, scrolling=True)
                except Exception as e:
                    st.error(f"تعذر قراءة أو إصدار الملف الجماعي: {e}")

    with tab_bg:
        st.subheader("إدارة الخلفيات")
        up = st.file_uploader("رفع خلفية جديدة PNG/JPG", type=["png","jpg","jpeg"], key="bg_upload")
        name_bg = st.text_input("اسم الخلفية بدون امتداد", "new_background")
        if st.button("حفظ الخلفية"):
            if up:
                ext = Path(up.name).suffix.lower()
                path = BG_DIR / f"{name_bg}{ext}"
                path.write_bytes(up.getbuffer())
                st.success("تم حفظ الخلفية")
                st.rerun()
            else:
                st.warning("اختر ملف خلفية أولاً.")
        st.write("الخلفيات الحالية:")
        for p in backgrounds():
            st.write("-", p.name)

    with tab_templates:
        st.subheader("القوالب")
        templates = sorted(TEMPLATE_DIR.glob("*.json"))
        if templates:
            chosen = st.selectbox("اختر قالباً", templates, format_func=lambda p: p.name)
            c1, c2 = st.columns(2)
            with c1:
                if st.button("تحميل القالب"):
                    load_template_file(chosen)
                    st.success("تم تحميل القالب")
                    st.rerun()
            with c2:
                st.download_button("تنزيل ملف القالب JSON", chosen.read_bytes(), file_name=chosen.name, mime="application/json")
        else:
            st.info("لا توجد قوالب محفوظة بعد.")

        st.divider()
        st.subheader("الخطوط")
        st.info("ضع ملفات الخطوط بصيغة TTF أو OTF داخل assets/fonts. النسخة V23 تختار خطاً عربياً آمناً تلقائياً عند النشر على Streamlit Cloud.")
        if FONT_DIR.exists():
            current_fonts = sorted([p.name for p in FONT_DIR.glob("*.ttf")] + [p.name for p in FONT_DIR.glob("*.otf")])
            if current_fonts:
                st.write("الخطوط الموجودة حالياً:")
                for f in current_fonts:
                    st.write("-", f)
            else:
                st.write("لا توجد خطوط داخل assets/fonts حالياً.")

        upload_template = st.file_uploader("استيراد قالب JSON", type=["json"])
        if upload_template and st.button("استيراد وتحميل القالب"):
            data = json.loads(upload_template.getvalue().decode("utf-8"))
            st.session_state.layout = data.get("layout", DEFAULT_LAYOUT)
            st.session_state.texts_template = data.get("texts_template", DEFAULT_TEXTS)
            st.success("تم استيراد القالب")
            st.rerun()

    with tab_log:
        st.subheader("سجل الشهادات")
        if EXCEL_FILE.exists():
            df = pd.read_excel(EXCEL_FILE)
            st.dataframe(df, use_container_width=True)
            st.download_button("تحميل السجل Excel", EXCEL_FILE.read_bytes(), file_name="issued_certificates.xlsx")
        else:
            st.info("لا يوجد سجل بعد.")

    with tab_users:
        users_management_tab()

    with tab_activity:
        activity_log_tab()


if __name__ == "__main__":
    main()
